
from __future__ import annotations

import io
import importlib.util
import logging
import os
import re
import shutil
from dataclasses import dataclass

logger = logging.getLogger(__name__)

CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150
MIN_CHUNK_LEN = 80
MAX_HEADING_LEN = 90
PDF_OCR_MIN_TEXT_CHARS = 50
PDF_OCR_RENDER_DPI = int(os.getenv("PDF_OCR_RENDER_DPI", "220"))
PDF_OCR_MAX_PAGES = int(os.getenv("PDF_OCR_MAX_PAGES", "120"))
PDF_OCR_LANG = os.getenv("PDF_OCR_LANG", "eng")

@dataclass
class PDFExtractionCandidate:
    parser: str
    text: str
    page_count: int
    non_empty_pages: int
    warnings: list[str]

def _clean_pdf_page_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def _strip_repeating_page_artifacts(page_texts: list[str]) -> list[str]:
    if len(page_texts) < 4:
        return page_texts

    threshold = max(3, int(len(page_texts) * 0.6))
    candidates: dict[str, int] = {}
    first_lines: list[str] = []
    last_lines: list[str] = []

    for page in page_texts:
        lines = [ln.strip() for ln in page.splitlines() if ln.strip()]
        if not lines:
            continue
        first_lines.append(lines[0])
        last_lines.append(lines[-1])

    def _count_lines(lines: list[str]):
        for line in lines:
            normalized = re.sub(r"\s+", " ", line).strip()
            if len(normalized) > 140:
                continue
            candidates[normalized] = candidates.get(normalized, 0) + 1

    _count_lines(first_lines)
    _count_lines(last_lines)
    to_remove = {line for line, count in candidates.items() if count >= threshold}
    if not to_remove:
        return page_texts

    cleaned_pages = []
    for page in page_texts:
        kept = []
        for line in page.splitlines():
            normalized = re.sub(r"\s+", " ", line).strip()
            if normalized and normalized in to_remove:
                continue
            kept.append(line)
        cleaned_pages.append("\n".join(kept).strip())
    return cleaned_pages

def _score_pdf_candidate(candidate: PDFExtractionCandidate) -> float:
    text = candidate.text or ""
    if not text.strip():
        return -1.0
    char_count = len(text)
    alpha_ratio = sum(ch.isalpha() for ch in text) / max(1, char_count)
    tokens = re.findall(r"[A-Za-z0-9]{2,}", text.lower())
    unique_ratio = (len(set(tokens)) / len(tokens)) if tokens else 0.0
    coverage = candidate.non_empty_pages / max(1, candidate.page_count)

    return (
        min(char_count, 400_000) * 0.55
        + alpha_ratio * 4_000
        + unique_ratio * 1_500
        + coverage * 8_000
    )

def _extract_with_pymupdf4llm(file_bytes: bytes) -> PDFExtractionCandidate | None:
    try:
        import fitz
        import pymupdf4llm
    except Exception:
        return None

    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_texts = [_clean_pdf_page_text(page.get_text("text", sort=True) or "") for page in doc]
        cleaned_pages = _strip_repeating_page_artifacts(page_texts)

        llm_text = _clean_pdf_page_text(pymupdf4llm.to_text(doc) or "")
        fallback_text = _clean_pdf_page_text("\n\n".join(cleaned_pages))
        final_text = llm_text if len(llm_text) >= len(fallback_text) * 0.7 else fallback_text

        return PDFExtractionCandidate(
            parser="pymupdf4llm",
            text=final_text,
            page_count=doc.page_count,
            non_empty_pages=sum(1 for p in cleaned_pages if p),
            warnings=[],
        )
    except Exception as e:
        logger.warning(f"pymupdf4llm extraction failed: {e}")
        return None
    finally:
        if doc is not None:
            doc.close()

def _extract_with_pymupdf(file_bytes: bytes) -> PDFExtractionCandidate | None:
    try:
        import fitz
    except Exception:
        return None

    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_texts: list[str] = []
        for page in doc:
            blocks = page.get_text("blocks", sort=True)
            block_lines = []
            for block in blocks:
                if len(block) > 4 and str(block[4]).strip():
                    block_lines.append(str(block[4]).strip())
            page_text = "\n".join(block_lines) if block_lines else (page.get_text("text", sort=True) or "")
            page_texts.append(_clean_pdf_page_text(page_text))

        cleaned_pages = _strip_repeating_page_artifacts(page_texts)
        return PDFExtractionCandidate(
            parser="pymupdf",
            text=_clean_pdf_page_text("\n\n".join(cleaned_pages)),
            page_count=doc.page_count,
            non_empty_pages=sum(1 for p in cleaned_pages if p),
            warnings=[],
        )
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed: {e}")
        return None
    finally:
        if doc is not None:
            doc.close()

def _extract_with_pdfplumber(file_bytes: bytes) -> PDFExtractionCandidate | None:
    try:
        import pdfplumber
    except Exception:
        return None

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_texts: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text(
                    x_tolerance=2,
                    y_tolerance=3,
                    layout=False,
                ) or ""
                page_texts.append(_clean_pdf_page_text(page_text))

            cleaned_pages = _strip_repeating_page_artifacts(page_texts)
            return PDFExtractionCandidate(
                parser="pdfplumber",
                text=_clean_pdf_page_text("\n\n".join(cleaned_pages)),
                page_count=len(pdf.pages),
                non_empty_pages=sum(1 for p in cleaned_pages if p),
                warnings=[],
            )
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
        return None

def _extract_with_pypdf2(file_bytes: bytes) -> PDFExtractionCandidate | None:
    try:
        import PyPDF2
    except Exception:
        return None

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        page_texts = [_clean_pdf_page_text((page.extract_text() or "")) for page in reader.pages]
        cleaned_pages = _strip_repeating_page_artifacts(page_texts)
        return PDFExtractionCandidate(
            parser="pypdf2",
            text=_clean_pdf_page_text("\n\n".join(cleaned_pages)),
            page_count=len(reader.pages),
            non_empty_pages=sum(1 for p in cleaned_pages if p),
            warnings=[],
        )
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
        return None

def _ocr_dependency_warnings() -> list[str]:
    warnings: list[str] = []
    if importlib.util.find_spec("pytesseract") is None:
        warnings.append("Missing OCR dependency: pytesseract")
    if shutil.which("tesseract") is None:
        warnings.append("Missing OCR system binary: tesseract")
    if importlib.util.find_spec("PIL") is None:
        warnings.append("Missing OCR dependency: Pillow")
    if importlib.util.find_spec("fitz") is None:
        warnings.append("Missing PDF renderer dependency: PyMuPDF")
    return warnings

def _extract_with_tesseract_ocr(file_bytes: bytes) -> PDFExtractionCandidate | None:
    dependency_warnings = _ocr_dependency_warnings()
    if dependency_warnings:
        page_count = 0
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = doc.page_count
            doc.close()
        except Exception:
            page_count = 0
        return PDFExtractionCandidate(
            parser="ocr-tesseract",
            text="",
            page_count=page_count,
            non_empty_pages=0,
            warnings=dependency_warnings,
        )

    try:
        import fitz
        import pytesseract
        from PIL import Image
    except Exception as e:
        logger.warning(f"OCR import failed: {e}")
        return PDFExtractionCandidate(
            parser="ocr-tesseract",
            text="",
            page_count=0,
            non_empty_pages=0,
            warnings=[f"OCR import failed: {e}"],
        )

    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        matrix = fitz.Matrix(PDF_OCR_RENDER_DPI / 72, PDF_OCR_RENDER_DPI / 72)
        page_texts: list[str] = []
        processed_pages = min(doc.page_count, PDF_OCR_MAX_PAGES)

        for page_index in range(processed_pages):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
            text = pytesseract.image_to_string(
                image,
                lang=PDF_OCR_LANG,
                config="--psm 6",
            )
            page_texts.append(_clean_pdf_page_text(text))

        cleaned_pages = _strip_repeating_page_artifacts(page_texts)
        warnings: list[str] = []
        if doc.page_count > processed_pages:
            warnings.append(
                f"OCR limited to first {processed_pages} pages; set PDF_OCR_MAX_PAGES to process more"
            )

        return PDFExtractionCandidate(
            parser="ocr-tesseract",
            text=_clean_pdf_page_text("\n\n".join(cleaned_pages)),
            page_count=doc.page_count,
            non_empty_pages=sum(1 for p in cleaned_pages if len(p) > 10),
            warnings=warnings,
        )
    except Exception as e:
        logger.warning(f"Tesseract OCR extraction failed: {e}")
        return PDFExtractionCandidate(
            parser="ocr-tesseract",
            text="",
            page_count=doc.page_count if doc is not None else 0,
            non_empty_pages=0,
            warnings=[f"OCR extraction failed: {e}"],
        )
    finally:
        if doc is not None:
            doc.close()

def extract_text_from_pdf_detailed(file_bytes: bytes) -> dict:
    candidates: list[PDFExtractionCandidate] = []
    warnings: list[str] = []

    for extractor in (
        _extract_with_pymupdf4llm,
        _extract_with_pymupdf,
        _extract_with_pdfplumber,
        _extract_with_pypdf2,
    ):
        candidate = extractor(file_bytes)
        if candidate is not None:
            candidates.append(candidate)

    best_text_len = max((len((candidate.text or "").strip()) for candidate in candidates), default=0)
    ocr_candidate: PDFExtractionCandidate | None = None
    if best_text_len < PDF_OCR_MIN_TEXT_CHARS:
        logger.info("PDF text extraction was empty or very short; trying OCR fallback")
        ocr_candidate = _extract_with_tesseract_ocr(file_bytes)
        if ocr_candidate is not None:
            if len((ocr_candidate.text or "").strip()) >= PDF_OCR_MIN_TEXT_CHARS:
                candidates.append(ocr_candidate)
            else:
                warnings.extend(ocr_candidate.warnings)

    if not candidates:
        missing = []
        if importlib.util.find_spec("fitz") is None:
            missing.append("PyMuPDF")
        if importlib.util.find_spec("pymupdf4llm") is None:
            missing.append("pymupdf4llm")
        if importlib.util.find_spec("pdfplumber") is None:
            missing.append("pdfplumber")
        if importlib.util.find_spec("PyPDF2") is None:
            missing.append("PyPDF2")
        if missing:
            warnings.append("Missing PDF parser dependencies: " + ", ".join(sorted(set(missing))))
        warnings.append("No PDF parser succeeded")
        return {
            "text": "",
            "parser": "",
            "page_count": 0,
            "non_empty_pages": 0,
            "warnings": warnings,
        }

    scored = sorted(
        ((candidate, _score_pdf_candidate(candidate)) for candidate in candidates),
        key=lambda item: item[1],
        reverse=True,
    )
    best, best_score = scored[0]
    logger.info(
        "PDF extraction selected parser=%s pages=%s non_empty_pages=%s score=%.2f",
        best.parser,
        best.page_count,
        best.non_empty_pages,
        best_score,
    )
    for candidate, score in scored[1:]:
        logger.debug("PDF extraction candidate parser=%s score=%.2f", candidate.parser, score)

    for c in candidates:
        warnings.extend(c.warnings)

    return {
        "text": best.text,
        "parser": best.parser,
        "page_count": best.page_count,
        "non_empty_pages": best.non_empty_pages,
        "warnings": sorted(set(warnings)),
    }

def extract_text_from_pdf(file_bytes: bytes) -> str:
    return extract_text_from_pdf_detailed(file_bytes).get("text", "")

def _extract_pages_with_pymupdf(file_bytes: bytes) -> list[dict] | None:
    try:
        import fitz
    except Exception:
        return None
    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            blocks = page.get_text("blocks", sort=True)
            block_lines = []
            for block in blocks:
                if len(block) > 4 and str(block[4]).strip():
                    block_lines.append(str(block[4]).strip())
            page_text = "\n".join(block_lines) if block_lines else (page.get_text("text", sort=True) or "")
            cleaned = _clean_pdf_page_text(page_text)
            pages.append({"page_num": i + 1, "text": cleaned, "char_count": len(cleaned)})
        return pages
    except Exception as e:
        logger.warning(f"PyMuPDF per-page extraction failed: {e}")
        return None
    finally:
        if doc is not None:
            doc.close()

def _extract_pages_with_pdfplumber(file_bytes: bytes) -> list[dict] | None:
    try:
        import pdfplumber
    except Exception:
        return None
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(x_tolerance=2, y_tolerance=3, layout=False) or ""
                cleaned = _clean_pdf_page_text(page_text)
                pages.append({"page_num": i + 1, "text": cleaned, "char_count": len(cleaned)})
            return pages
    except Exception as e:
        logger.warning(f"pdfplumber per-page extraction failed: {e}")
        return None

def _extract_pages_with_pypdf2(file_bytes: bytes) -> list[dict] | None:
    try:
        import PyPDF2
    except Exception:
        return None
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            cleaned = _clean_pdf_page_text(page.extract_text() or "")
            pages.append({"page_num": i + 1, "text": cleaned, "char_count": len(cleaned)})
        return pages
    except Exception as e:
        logger.warning(f"PyPDF2 per-page extraction failed: {e}")
        return None

def _extract_pages_with_tesseract_ocr(file_bytes: bytes) -> list[dict] | None:
    dependency_warnings = _ocr_dependency_warnings()
    if dependency_warnings:
        return None

    try:
        import fitz
        import pytesseract
        from PIL import Image
    except Exception as e:
        logger.warning(f"OCR per-page import failed: {e}")
        return None

    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        matrix = fitz.Matrix(PDF_OCR_RENDER_DPI / 72, PDF_OCR_RENDER_DPI / 72)
        processed_pages = min(doc.page_count, PDF_OCR_MAX_PAGES)
        pages: list[dict] = []

        for page_index in range(processed_pages):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("L")
            text = pytesseract.image_to_string(
                image,
                lang=PDF_OCR_LANG,
                config="--psm 6",
            )
            cleaned = _clean_pdf_page_text(text)
            pages.append({
                "page_num": page_index + 1,
                "text": cleaned,
                "char_count": len(cleaned),
                "parser": "ocr-tesseract",
            })

        if sum(p["char_count"] for p in pages) < PDF_OCR_MIN_TEXT_CHARS:
            return None
        return pages
    except Exception as e:
        logger.warning(f"Tesseract OCR per-page extraction failed: {e}")
        return None
    finally:
        if doc is not None:
            doc.close()

def _score_pages(pages: list[dict]) -> float:
    if not pages:
        return -1.0
    total_chars = sum(p["char_count"] for p in pages)
    non_empty = sum(1 for p in pages if p["char_count"] > 10)
    coverage = non_empty / max(1, len(pages))
    all_text = " ".join(p["text"] for p in pages if p["text"])
    alpha_ratio = (sum(c.isalpha() for c in all_text) / max(1, len(all_text))) if all_text else 0.0
    return min(total_chars, 400_000) * 0.55 + alpha_ratio * 4_000 + coverage * 8_000

def extract_pages_from_pdf(file_bytes: bytes) -> list[dict]:
    candidates: list[tuple[str, list[dict]]] = []
    for name, extractor in [
        ("pymupdf", _extract_pages_with_pymupdf),
        ("pdfplumber", _extract_pages_with_pdfplumber),
        ("pypdf2", _extract_pages_with_pypdf2),
    ]:
        result = extractor(file_bytes)
        if result:
            candidates.append((name, result))

    if not candidates:
        logger.warning("extract_pages_from_pdf: no parser succeeded")
        return []

    best_name, best_pages = max(candidates, key=lambda x: _score_pages(x[1]))
    score = _score_pages(best_pages)
    total_chars = sum(p["char_count"] for p in best_pages)
    if total_chars < PDF_OCR_MIN_TEXT_CHARS:
        logger.info("Per-page PDF extraction was empty or very short; trying OCR fallback")
        ocr_pages = _extract_pages_with_tesseract_ocr(file_bytes)
        if ocr_pages:
            best_name = "ocr-tesseract"
            best_pages = ocr_pages
            score = _score_pages(best_pages)
        else:
            return []

    logger.info(
        "Per-page extraction selected parser=%s pages=%d score=%.1f",
        best_name, len(best_pages), score,
    )

    page_texts = [p["text"] for p in best_pages]
    cleaned_texts = _strip_repeating_page_artifacts(page_texts)
    result = []
    for i, p in enumerate(best_pages):
        text = cleaned_texts[i]
        result.append({
            "page_num": p["page_num"],
            "text": text,
            "char_count": len(text),
            "parser": best_name,
        })
    return result

def chunk_pages_with_tracking(
    pages: list[dict],
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
    min_length: int = MIN_CHUNK_LEN,
) -> list[dict]:
    result: list[dict] = []
    pending_text = ""
    pending_page_start: int | None = None

    for page_info in pages:
        page_num: int = page_info["page_num"]
        page_text: str = page_info.get("text", "")

        if not page_text or not page_text.strip():
            continue

        if pending_text:
            combined_text = pending_text + "\n\n" + page_text
            page_start = pending_page_start
        else:
            combined_text = page_text
            page_start = page_num

        normalized = _normalize_text(combined_text)
        if not normalized or len(normalized) < min_length:
            pending_text = combined_text
            pending_page_start = page_start
            continue

        chunks = _chunk_normalized_text(normalized, chunk_size, overlap, min_length)
        if not chunks:
            pending_text = combined_text
            pending_page_start = page_start
            continue

        page_end = page_num
        page_label = str(page_start) if page_start == page_end else f"{page_start}-{page_end}"
        for chunk_text_item in chunks:
            result.append({
                "text": chunk_text_item,
                "page_start": page_start,
                "page_end": page_end,
                "page_label": page_label,
            })

        pending_text = ""
        pending_page_start = None

    if pending_text:
        normalized = _normalize_text(pending_text)
        if normalized and len(normalized) >= min_length:
            chunks = _chunk_normalized_text(normalized, chunk_size, overlap, min_length)
            last_page = pages[-1]["page_num"] if pages else pending_page_start
            page_label = (
                str(pending_page_start)
                if pending_page_start == last_page
                else f"{pending_page_start}-{last_page}"
            )
            for chunk_text_item in chunks:
                result.append({
                    "text": chunk_text_item,
                    "page_start": pending_page_start,
                    "page_end": last_page,
                    "page_label": page_label,
                })

    return result

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace").strip()

_BULLET_RE = re.compile("^\\s*(?:[-*]|\\u2022|\\d+[\\.\\)])\\s+")
_HEADING_RE = re.compile(
    r"^(chapter|unit|module|lesson|section|part)\s+([0-9ivxlcdm]+)([:\-\.\s].*)?$",
    re.IGNORECASE,
)
_ALLCAPS_RE = re.compile(r"^[A-Z0-9][A-Z0-9\s\-:]{6,}$")

def _normalize_text(text: str) -> str:
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    lines = [ln.strip() for ln in text.split("\n")]
    paragraphs: list[str] = []
    current: list[str] = []

    def _flush_current():
        if current:
            paragraphs.append(" ".join(current).strip())
            current.clear()

    for line in lines:
        if not line:
            _flush_current()
            continue

        if _BULLET_RE.match(line):
            _flush_current()
            paragraphs.append(line.strip())
            continue

        current.append(line)

    _flush_current()

    cleaned = [re.sub(r"\s+", " ", p).strip() for p in paragraphs if p.strip()]
    return "\n\n".join(cleaned)

def _is_heading(line: str) -> bool:
    if not line:
        return False
    if len(line) > MAX_HEADING_LEN:
        return False
    compact = re.sub(r"\s+", " ", line.strip())
    if len(compact) < 4:
        return False
    if _HEADING_RE.match(compact):
        return True
    if _ALLCAPS_RE.match(compact) and len(compact.split()) <= 8:
        return True
    return False

def extract_chapter_headings(text: str, limit: int = 16) -> list[str]:
    if not text:
        return []
    headings: list[str] = []
    seen = set()
    for line in text.splitlines():
        candidate = line.strip()
        if not candidate:
            continue
        if _is_heading(candidate):
            normalized = re.sub(r"\s+", " ", candidate)
            if normalized not in seen:
                headings.append(normalized)
                seen.add(normalized)
        if len(headings) >= limit:
            break
    return headings

_GRADE_PATTERNS: list[tuple[str, re.Pattern]] = [
    ("AP", re.compile(r"\bAP\b|\bAdvanced Placement\b", re.IGNORECASE)),
    ("Honors", re.compile(r"\bHonors?\b", re.IGNORECASE)),
    ("Grade 12", re.compile(r"\b(12th|grade\s*12|grade\s*twelve|senior)\b", re.IGNORECASE)),
    ("Grade 11", re.compile(r"\b(11th|grade\s*11|grade\s*eleven|junior)\b", re.IGNORECASE)),
    ("Grade 10", re.compile(r"\b(10th|grade\s*10|grade\s*ten|sophomore)\b", re.IGNORECASE)),
    ("Grade 9", re.compile(r"\b(9th|grade\s*9|grade\s*nine|freshman)\b", re.IGNORECASE)),
]

def infer_grade_level(text: str) -> str:
    if not text:
        return ""
    for label, pattern in _GRADE_PATTERNS:
        if pattern.search(text):
            return label
    return ""

def _tail_overlap(text: str, overlap: int) -> str:
    if overlap <= 0 or not text:
        return ""
    if len(text) <= overlap:
        return text.strip()
    snippet = text[-overlap:]
    idx = snippet.find(" ")
    if idx != -1 and idx + 1 < len(snippet):
        snippet = snippet[idx + 1 :]
    return snippet.strip()

def _sliding_window_chunks(
    text: str,
    chunk_size: int,
    overlap: int,
    min_length: int,
) -> list[str]:
    if not text or len(text) < min_length:
        return []

    step = chunk_size - overlap
    chunks: list[str] = []
    start = 0

    while start < len(text):
        chunk = text[start : start + chunk_size].strip()
        if len(chunk) >= min_length:
            chunks.append(chunk)
        start += step

    return chunks

def _chunk_normalized_text(
    normalized: str,
    chunk_size: int,
    overlap: int,
    min_length: int,
) -> list[str]:
    if not normalized or len(normalized) < min_length:
        return []

    paragraphs = [p for p in normalized.split("\n\n") if p.strip()]
    if len(paragraphs) <= 1:
        return _sliding_window_chunks(normalized, chunk_size, overlap, min_length)

    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(para) > chunk_size:
            if len(current) >= min_length:
                chunks.append(current.strip())
            current = ""
            chunks.extend(_sliding_window_chunks(para, chunk_size, overlap, min_length))
            continue

        if not current:
            current = para
            continue

        if len(current) + 2 + len(para) <= chunk_size:
            current = f"{current}\n\n{para}"
        else:
            if len(current) >= min_length:
                chunks.append(current.strip())
                overlap_text = _tail_overlap(current, overlap)
                if overlap_text and len(overlap_text) + 2 + len(para) <= chunk_size:
                    current = f"{overlap_text}\n\n{para}"
                else:
                    current = para
            else:
                # current is too short to stand alone as a chunk -- merge it
                # forward instead of overwriting it with para, or this text
                # is silently dropped from the document's embedded chunks.
                current = f"{current}\n\n{para}"

    if current and len(current) >= min_length:
        chunks.append(current.strip())

    return chunks

def _split_sections(text: str) -> list[tuple[str, str]]:
    if not text:
        return []
    sections: list[tuple[str, str]] = []
    current_heading = "Intro"
    buffer: list[str] = []

    def _flush():
        if buffer:
            sections.append((current_heading, "\n".join(buffer).strip()))
            buffer.clear()

    for line in text.splitlines():
        if _is_heading(line):
            _flush()
            current_heading = re.sub(r"\s+", " ", line.strip())
        else:
            buffer.append(line)

    _flush()
    return sections

def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
    min_length: int = MIN_CHUNK_LEN,
    toc_aware: bool = False,
) -> list[str]:
    if not text or len(text) < min_length:
        return []

    normalized = _normalize_text(text)
    if not normalized or len(normalized) < min_length:
        return []

    if not toc_aware:
        return _chunk_normalized_text(normalized, chunk_size, overlap, min_length)

    sections = _split_sections(text)
    if len(sections) <= 1:
        return _chunk_normalized_text(normalized, chunk_size, overlap, min_length)

    chunks: list[str] = []
    for heading, section_text in sections:
        if not section_text:
            continue
        normalized_section = _normalize_text(section_text)
        if not normalized_section:
            continue

        prefix = "" if heading.lower() == "intro" else heading
        effective_size = chunk_size
        if prefix:
            effective_size = max(min_length, chunk_size - len(prefix) - 2)

        section_chunks = _chunk_normalized_text(
            normalized_section,
            effective_size,
            overlap,
            min_length,
        )
        if prefix:
            section_chunks = [f"{prefix}\n\n{chunk}" for chunk in section_chunks]
        chunks.extend(section_chunks)

    return chunks

def process_upload(
    file_bytes: bytes,
    filename: str,
    subject: str = "",
    grade_level: str = "",
    scope: str = "private",
    source_url: str = "",
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    toc_aware: bool = True,
) -> dict:
    lower_name = (filename or "").lower()
    text = ""
    error = None
    detected_subject = ""
    detected_grade = ""
    chapters: list[str] = []
    pdf_parser = ""
    pdf_page_count = 0
    pdf_non_empty_pages = 0
    extraction_warnings: list[str] = []
    _pdf_pages: list[dict] = []

    try:
        if lower_name.endswith(".pdf"):
            _pdf_pages = extract_pages_from_pdf(file_bytes)
            if _pdf_pages:
                pdf_parser = _pdf_pages[0].get("parser", "pymupdf")
                pdf_page_count = len(_pdf_pages)
                pdf_non_empty_pages = sum(1 for p in _pdf_pages if p["char_count"] > 10)
                text = "\n\n".join(p["text"] for p in _pdf_pages if p["text"])
            else:
                extracted = extract_text_from_pdf_detailed(file_bytes)
                text = extracted.get("text", "")
                pdf_parser = extracted.get("parser", "")
                pdf_page_count = extracted.get("page_count", 0) or 0
                pdf_non_empty_pages = extracted.get("non_empty_pages", 0) or 0
                extraction_warnings = extracted.get("warnings", []) or []
        elif lower_name.endswith(".docx"):
            try:
                from docx import Document
            except ImportError as exc:
                raise RuntimeError("DOCX support is unavailable on the server") from exc
            document = Document(io.BytesIO(file_bytes))
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
            text = "\n\n".join(blocks)
        elif lower_name.endswith((".txt", ".md")):
            text = extract_text_from_txt(file_bytes)
        else:
            error = (
                f"Unsupported file type: '{filename}'. "
                "Please upload a .pdf, .docx, .txt, or .md file."
            )
    except Exception as e:
        error = f"Text extraction failed: {e}"
        logger.error(error)

    if not text and not error:
        if lower_name.endswith(".pdf") and any(
            w.lower().startswith("missing pdf parser dependencies")
            for w in extraction_warnings
        ):
            error = "PDF parsing dependencies are missing on the server."
        elif lower_name.endswith(".pdf") and any(
            "ocr" in w.lower() or "tesseract" in w.lower()
            for w in extraction_warnings
        ):
            error = (
                "No extractable text found in this scanned or image-only PDF, "
                "and OCR is not available on the server. Install Tesseract OCR "
                "and the pytesseract Python package, then retry the upload."
            )
        elif lower_name.endswith(".pdf") and pdf_page_count > 0:
            error = (
                "No extractable text found in this PDF. "
                "It appears to be image-only or scanned content."
            )
        else:
            error = "No text could be extracted from this file."

    if text:
        chapters = extract_chapter_headings(text)
        if not subject:
            try:
                from services import context_store
                detected_subject = context_store.infer_subject(f"{filename} {text[:4000]}", default="")
                subject = detected_subject or subject
            except Exception:
                detected_subject = ""
        if not grade_level:
            detected_grade = infer_grade_level(f"{filename}\n{text[:4000]}")
            grade_level = detected_grade or grade_level

    chunk_dicts: list[dict] = []
    if _pdf_pages and text:
        chunk_dicts = chunk_pages_with_tracking(
            _pdf_pages,
            chunk_size=chunk_size,
            overlap=chunk_overlap,
            min_length=MIN_CHUNK_LEN,
        )
    elif text:
        plain_chunks = chunk_text(
            text,
            chunk_size=chunk_size,
            overlap=chunk_overlap,
            min_length=MIN_CHUNK_LEN,
            toc_aware=toc_aware,
        )
        chunk_dicts = [{"text": c, "page_start": None, "page_end": None, "page_label": ""} for c in plain_chunks]

    chunks: list[str] = [c["text"] for c in chunk_dicts]
    chunk_pages: list[dict] = [
        {"page_start": c["page_start"], "page_end": c["page_end"], "page_label": c["page_label"]}
        for c in chunk_dicts
    ]

    return {
        "chunks":      chunks,
        "chunk_pages": chunk_pages,
        "chunk_count": len(chunks),
        "char_count":  len(text),
        "filename":    filename,
        "subject":     subject,
        "grade_level": grade_level,
        "scope":       scope,
        "error":       error,
        "detected_subject": detected_subject,
        "detected_grade": detected_grade,
        "chapters":    chapters,
        "chunk_size":  chunk_size,
        "chunk_overlap": chunk_overlap,
        "toc_aware":   toc_aware,
        "pdf_parser": pdf_parser,
        "pdf_page_count": pdf_page_count,
        "pdf_non_empty_pages": pdf_non_empty_pages,
        "extraction_warnings": extraction_warnings,
        "has_page_tracking": bool(_pdf_pages),
    }
